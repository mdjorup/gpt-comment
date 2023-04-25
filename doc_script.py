
from docx import Document

from comments import Comment, QuotedComment

essay = 'My proudest achievement was when my classmate and I succeeded in creating a functional bridge that was able to hold up to nine pounds of sand. For context, in seventh grade, I chose engineering as my elective because I wanted to create things that benefitted society. I was assigned to create a miniature bridge with a classmate out of wood strips and glue. I created a blueprint based on other designs and cooperated with my classmate by setting goals, assigning roles, etc.\nOur bridge ended up being able to hold 9 pounds of sand. This experience helped me learn that being able to collaborate with others is a skill that will be needed often throughout life. Knowing how to create goals and consistently work towards them aids me in achieving other long-term goals in my life. Despite minor setbacks, I remained determined and persistent in my efforts. I hope to be able to apply these skills to activities in TJ such as the Innovations in Concepts of Engineering (ICE).'

coms : list[QuotedComment]= []

coms.append(QuotedComment(comment='Replace with "benefited"', quote='benefitted', start_index=248, length=10))
coms.append(QuotedComment(comment='Add a comma after "designs"', quote='I created a blueprint based on other designs', start_index=358, length=44))
coms.append(QuotedComment(comment='Replace "etc." with "and"', quote='cooperated with my classmate by setting goals, assigning roles, etc.', start_index=407, length=68))
coms.append(QuotedComment(comment='Add a comma after "bridge"', quote='Our bridge ended up', start_index=476, length=19))
coms.append(QuotedComment(comment='Add a comma after "others"', quote='being able to collaborate with others is a skill that will be needed often throughout life', start_index=570, length=90))
coms.append(QuotedComment(comment='Replace "aids" with "will aid"', quote='Knowing how to create goals and consistently work towards them aids me in achieving other long-term goals in my life.', start_index=662, length=117))
coms.append(QuotedComment(comment='Add a comma after "setbacks"', quote='Despite minor setbacks, I remained determined and persistent in my efforts.', start_index=780, length=75))
coms.append(QuotedComment(comment='Add a comma after "TJ" and capitalize "innovations" and "concepts" to make it "Innovations in Concepts of Engineering (ICE)."', quote='I hope to be able to apply these skills to activities in TJ such as the Innovations in Concepts of Engineering (ICE).', start_index=856, length=117))
coms.append(QuotedComment(comment='It is always good to give credit where credit is due. Consider citing your sources to show that you researched and used existing designs in your work.', quote='I created a blueprint based on other designs', start_index=358, length=44))
coms.append(QuotedComment(comment='Instead of just stating the result, consider adding more details about the process of testing the bridge\'s weight capacity. This will give readers a better understanding of your achievement.', quote='Our bridge ended up being able to hold 9 pounds of sand.', start_index=476, length=56))
coms.append(QuotedComment(comment='This is a good point, but it would be stronger if you could provide specific examples of other long-term goals you have set for yourself and how you have worked towards achieving them.', quote='Knowing how to create goals and consistently work towards them aids me in achieving other long-term goals in my life.', start_index=662, length=117))
coms.append(QuotedComment(comment='This is a great example of resilience, but it would be even better if you could provide specific examples of the setbacks you faced and how you overcame them.', quote='Despite minor setbacks, I remained determined and persistent in my efforts.', start_index=780, length=75))
coms.append(QuotedComment(comment='This is a good way to tie your achievement to your interest in attending the school, but it would be stronger if you could provide more detail about what ICE is and how your skills would be useful in that program.', quote='I hope to be able to apply these skills to activities in TJ such as the Innovations in Concepts of Engineering (ICE).', start_index=856, length=117))
gc = Comment(comment='This essay is a good start, but there are some areas for improvement. Firstly, it would be better to provide more specific details about the adversity faced during the bridge-building project. What were the setbacks and how did you overcome them? This will demonstrate resilience and determination, which are key qualities the prompt is looking for. Additionally, it would be good to expand on how this achievement relates to your long-term goals and aspirations. How will the skills you learned in this project help you succeed in your future endeavors? Finally, it would be beneficial to incorporate more personal reflection and insight. What did you learn about yourself through this experience? How has it shaped your perspective on adversity and perseverance? By addressing these areas, you can provide a more comprehensive and compelling essay.', prefix="General Comment:")


coms.sort(key= lambda x: x.start_index)
# for i, com in enumerate(coms):
#     print(com.class_repr())

# quit() 

doc = Document()

paragraph = doc.add_paragraph()

current_index = 0

for i, comment in enumerate(coms):

    start = comment.start_index
    end = min(start + comment.length, len(essay))
    if i < len(coms) - 1:
        end = min(end, coms[i+1].start_index)
        coms[i+1].start_index = max(coms[i+1].start_index, end)
    
    if start > current_index:
        paragraph.add_run(essay[current_index:start])

    
    new_run = paragraph.add_run(essay[start:end])
    new_run.add_comment(comment.comment)

    current_index = end

if current_index < len(essay):
    paragraph.add_run(essay[current_index:])

new_paragraph = doc.add_paragraph("\n" + str(gc))

doc.save('test.docx')





