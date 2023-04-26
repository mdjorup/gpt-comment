import asyncio
import time
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from tqdm import tqdm

from config import config
from essay import PSEEssay, SPSEssay
from services.drive_service import DriveService
from services.sheets_service import SheetsService
from utils import float_to_dollar


class StudentEntry:
    
    def __init__(self,  essays : dict, 
                        row_index : int, 
                        student_email : str, 
                        parent_email : str,
                        gpa : float,
                        middle_school : str):
        
        self.sps_essays : list[SPSEssay] = []
        self.pse_essays : list[PSEEssay] = []

        for key, value in essays.items():
            if value.strip() == "":
                continue

            if key.startswith(config["spreadsheet"]["sps_start_indicator"]):
                prompt = key[4:]
                essay = value
                sps_essay = SPSEssay(prompt, essay)
                self.sps_essays.append(sps_essay)
            elif key.startswith(config["spreadsheet"]["pse_start_indicator"]):
                prompt = key[4:]
                essay = value
                pse_essay = PSEEssay(prompt, essay)
                self.pse_essays.append(pse_essay)
            else:
                continue



        self.row_index = row_index 
        self.student_email = student_email
        self.parent_email = parent_email
        self.gpa : float = gpa
        self.middle_school = middle_school
        


    def __str__(self):
        return f"StudentEntry(row_index={self.row_index}, student_email={self.student_email}, n_sps_essays={len(self.sps_essays)}, n_pse_essays={len(self.pse_essays)})"


    def update_completed(self, document_link : str):
        ss = SheetsService()
        completed_column = config["spreadsheet"]["completed_column"]
        document_link_column = config["spreadsheet"]["document_link_column"]
        ss.update_cell(self.row_index, completed_column, True)
        ss.update_cell(self.row_index, document_link_column, document_link)


    
    async def process(self):
        print("Processing Submission for " + self.student_email)

        ### PART 2: PROCESS STUDENT ENTRIES - everything to do with analyzing the essays
        t0 = time.time()
        async with asyncio.TaskGroup() as tg:
            sps_essays_bar = tqdm(self.sps_essays, desc='SPS Essays')
            for essay in self.sps_essays:
                tg.create_task(essay.process(sps_essays_bar))
                
                time.sleep(2)
            
            pse_essays_bar =  tqdm(self.pse_essays, desc='PSE Essays')
            for essay in self.pse_essays:
                tg.create_task(essay.process(pse_essays_bar))
                time.sleep(2)
        t1 = time.time() 

        ### PART 3: GENERATE REPORTS - create word doc, write entries

        total_cost = 0
        document = Document()
        title = document.add_paragraph()
        title_run = title.add_run("EduAvenues/TJTestPrep Essay Feedback")
        title_run.font.color.rgb = RGBColor.from_string("018AFD")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = 1 # centers the title
        
        se_paragraph = document.add_paragraph()
        se_run = se_paragraph.add_run("Student Email: ")
        se_paragraph.add_run(self.student_email)
        se_run.font.bold = True

        pe_paragraph = document.add_paragraph()
        pe_run = pe_paragraph.add_run("Parent Email: ")
        pe_paragraph.add_run(self.parent_email)
        pe_run.font.bold = True

        ms_paragraph = document.add_paragraph()
        ms_run = ms_paragraph.add_run("Middle School: ")
        ms_paragraph.add_run(self.middle_school)
        ms_run.font.bold = True

        gpa_paragraph = document.add_paragraph()
        gpa_run = gpa_paragraph.add_run("GPA: ")
        gpa_paragraph.add_run(str(round(self.gpa, 2)))
        gpa_run.font.bold = True

        date_paragraph = document.add_paragraph()
        date_run = date_paragraph.add_run("Date: ")
        date_paragraph.add_run(datetime.now().strftime("%m/%d/%Y"))
        date_run.font.bold = True

        essay_count = 1

        for essay in self.sps_essays:
            prompt_paragraph = document.add_paragraph()
            prompt_run = prompt_paragraph.add_run(f"Prompt {essay_count}: " + essay.prompt)
            prompt_run.font.bold = True

            essay.add_to_doc(document)
            essay_count += 1
            total_cost += essay.processing_costs
        
        for essay in self.pse_essays:
            document.add_paragraph(f"Prompt {essay_count}: " + essay.prompt)
            essay.add_to_doc(document)
            essay_count += 1
            total_cost += essay.processing_costs


        print("Submission Processing Cost: ", float_to_dollar(total_cost))
        print("Time taken: ", t1 - t0, "seconds")
        print("\n")

        document_name = config["generated_document_name"] + "-" + self.student_email + ".docx"

        footer = document.sections[0].footer

        paragraph = footer.paragraphs[0]

        left_part = paragraph.add_run('© EduAvenues LLC\t\t')
        left_part.font.color.rgb = RGBColor.from_string("018AFD")

        # Set the right side of the footer
        right_part = paragraph.add_run('EduAvenues')
        right_part.font.color.rgb = RGBColor.from_string("018AFD")
        right_part.font.name = "Avenir"
        right_part.font.size = Pt(14)

        # This code saves the word document locally
        #document.save(document_name)


        ### PART 4: UPDATE SHEET & UPLOAD

        ds = DriveService()

        link = ds.upload_word_doc(document_name, document)

        self.update_completed(link)

        return total_cost
        
        

        
        